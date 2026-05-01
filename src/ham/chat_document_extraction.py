"""
Server-side document text extraction for workspace chat attachments.

Extracted text is used only when assembling LLM message content
(``to_llm_message_content``); it must not be persisted in session storage,
returned in attachment APIs, or logged in full.
"""
from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO, StringIO
from typing import Literal

MAX_EXTRACTED_CHARS_PER_FILE = 30_000
MAX_EXTRACTED_CHARS_PER_MESSAGE = 60_000
MAX_PDF_PAGES = 25
MAX_DOCX_PARAGRAPHS = 2_000
MAX_DOCX_TABLES = 200
MAX_SPREADSHEET_SHEETS = 5
MAX_SPREADSHEET_ROWS = 100
MAX_SPREADSHEET_COLS = 30

ExtractionStatus = Literal["extracted", "truncated", "unsupported", "failed", "empty"]


def _is_spreadsheet_mime(m: str) -> bool:
    x = (m or "").strip().lower()
    return x in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "text/csv",
        "application/csv",
    )


def _is_legacy_xls_placeholder(filename: str, mime: str) -> bool:
    fn = (filename or "").lower()
    m = (mime or "").strip().lower()
    return fn.endswith(".xls") or m == "application/vnd.ms-excel"


def _cell_str(v: object) -> str:
    if v is None:
        return ""
    return str(v).replace("\n", " ").replace("\r", " ").strip()


@dataclass(frozen=True)
class DocumentExtractionResult:
    filename: str
    mime: str
    status: ExtractionStatus
    text: str
    truncated: bool
    error_reason: str | None


def _cap_body(body: str, limit: int) -> tuple[str, bool]:
    if len(body) <= limit:
        return body, False
    return body[:limit], True


def _decode_plain_text(raw: bytes) -> str:
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("utf-8", errors="replace")


def _extract_xlsx(raw: bytes) -> tuple[str, ExtractionStatus, bool, str | None]:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "", "failed", False, "xlsx_library_unavailable"

    try:
        wb = load_workbook(filename=BytesIO(raw), read_only=True, data_only=True)
    except Exception:
        return "", "failed", False, "invalid_xlsx"

    parts: list[str] = []
    truncated = False
    sheet_names_all = list(wb.sheetnames)
    sheet_names = sheet_names_all[:MAX_SPREADSHEET_SHEETS]
    if len(sheet_names_all) > MAX_SPREADSHEET_SHEETS:
        truncated = True

    for sname in sheet_names:
        ws = wb[sname]
        rows_buf: list[list[str]] = []
        row_idx = 0
        try:
            for row in ws.iter_rows(
                min_row=1,
                max_row=MAX_SPREADSHEET_ROWS,
                max_col=MAX_SPREADSHEET_COLS,
                values_only=True,
            ):
                row_idx += 1
                cells = [_cell_str(c) for c in (row or ())]
                while len(cells) < MAX_SPREADSHEET_COLS:
                    cells.append("")
                rows_buf.append(cells[:MAX_SPREADSHEET_COLS])
        except Exception:
            wb.close()
            return "", "failed", False, "invalid_xlsx"

        if row_idx >= MAX_SPREADSHEET_ROWS:
            truncated = True

        block: list[str] = [f"Sheet: {sname}"]
        if rows_buf:
            hdr = rows_buf[0]
            block.append(f"Columns: {' | '.join(hdr)}")
            block.append("Rows:")
            for data_row in rows_buf[1:]:
                block.append(" | ".join(data_row))
        parts.append("\n".join(block))

    try:
        wb.close()
    except Exception:
        pass

    merged = "\n\n".join(parts).strip()
    if not merged:
        return "", "empty", False, None
    capped, cut = _cap_body(merged, MAX_EXTRACTED_CHARS_PER_FILE)
    st: ExtractionStatus = "truncated" if (cut or truncated) else "extracted"
    return capped, st, cut or truncated, None


def _extract_csv(raw: bytes) -> tuple[str, ExtractionStatus, bool, str | None]:
    import csv

    text = _decode_plain_text(raw)
    rows: list[list[str]] = []
    truncated = False
    try:
        reader = csv.reader(StringIO(text))
        for i, row in enumerate(reader):
            if i >= MAX_SPREADSHEET_ROWS:
                truncated = True
                break
            rows.append([_cell_str(c) for c in row[:MAX_SPREADSHEET_COLS]])
            if len(row) > MAX_SPREADSHEET_COLS:
                truncated = True
    except Exception:
        return "", "failed", False, "invalid_csv"

    if not rows:
        return "", "empty", False, None

    block: list[str] = []
    block.append(f"Columns: {' | '.join(rows[0])}")
    block.append("Rows:")
    for data_row in rows[1:]:
        block.append(" | ".join(data_row))
    merged = "\n".join(block).strip()
    if not merged:
        return "", "empty", False, None
    capped, cut = _cap_body(merged, MAX_EXTRACTED_CHARS_PER_FILE)
    st: ExtractionStatus = "truncated" if (cut or truncated) else "extracted"
    return capped, st, cut or truncated, None


def _extract_pdf(raw: bytes) -> tuple[str, ExtractionStatus, bool, str | None]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "", "failed", False, "pdf_library_unavailable"

    try:
        reader = PdfReader(BytesIO(raw))
    except Exception:
        return "", "failed", False, "invalid_pdf"

    parts: list[str] = []
    page_count = min(len(reader.pages), MAX_PDF_PAGES)
    truncated_pages = len(reader.pages) > MAX_PDF_PAGES
    for i in range(page_count):
        try:
            page = reader.pages[i]
            t = (page.extract_text() or "").strip()
        except Exception:
            continue
        if t:
            parts.append(t)
    merged = "\n\n".join(parts).strip()
    if not merged:
        return "", "empty", False, None
    capped, cut = _cap_body(merged, MAX_EXTRACTED_CHARS_PER_FILE)
    st: ExtractionStatus
    if cut or truncated_pages:
        st = "truncated"
    else:
        st = "extracted"
    return capped, st, cut or truncated_pages, None


def _extract_docx(raw: bytes) -> tuple[str, ExtractionStatus, bool, str | None]:
    try:
        from docx import Document
    except ImportError:
        return "", "failed", False, "docx_library_unavailable"

    try:
        doc = Document(BytesIO(raw))
    except Exception:
        return "", "failed", False, "invalid_docx"

    lines: list[str] = []
    char_budget = MAX_EXTRACTED_CHARS_PER_FILE
    truncated = False

    for pi, p in enumerate(doc.paragraphs):
        if pi >= MAX_DOCX_PARAGRAPHS:
            truncated = True
            break
        t = (p.text or "").strip()
        if not t:
            continue
        if len(t) > char_budget:
            lines.append(t[:char_budget])
            char_budget = 0
            truncated = True
            break
        lines.append(t)
        char_budget -= len(t) + 1
        if char_budget <= 0:
            truncated = True
            break

    if char_budget > 0:
        for ti, table in enumerate(doc.tables):
            if ti >= MAX_DOCX_TABLES:
                truncated = True
                break
            for row in table.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if not t:
                        continue
                    if len(t) > char_budget:
                        lines.append(t[:char_budget])
                        char_budget = 0
                        truncated = True
                        break
                    lines.append(t)
                    char_budget -= len(t) + 1
                    if char_budget <= 0:
                        truncated = True
                        break
                if char_budget <= 0:
                    break
            if char_budget <= 0:
                break

    merged = "\n".join(lines).strip()
    if not merged:
        return "", "empty", False, None
    capped, cut = _cap_body(merged, MAX_EXTRACTED_CHARS_PER_FILE)
    st: ExtractionStatus = "truncated" if (cut or truncated) else "extracted"
    return capped, st, cut or truncated, None


def extract_document_bytes(*, filename: str, mime: str, raw: bytes) -> DocumentExtractionResult:
    """
    Best-effort extraction. Never raises; failures become ``failed`` / ``unsupported`` results.
    """
    fn = (filename or "").strip() or "attachment"
    m = (mime or "").strip().lower()
    if m == "image/jpg":
        m = "image/jpeg"

    if m in ("text/plain", "text/markdown"):
        body = _decode_plain_text(raw)
        capped, cut = _cap_body(body, MAX_EXTRACTED_CHARS_PER_FILE)
        if not capped.strip():
            return DocumentExtractionResult(
                filename=fn,
                mime=m,
                status="empty",
                text="",
                truncated=False,
                error_reason=None,
            )
        st: ExtractionStatus = "truncated" if cut else "extracted"
        return DocumentExtractionResult(
            filename=fn,
            mime=m,
            status=st,
            text=capped,
            truncated=cut,
            error_reason=None,
        )

    if m == "application/pdf":
        text, st, trunc, err = _extract_pdf(raw)
        return DocumentExtractionResult(
            filename=fn,
            mime=m,
            status=st,
            text=text,
            truncated=trunc,
            error_reason=err,
        )

    if m == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text, st, trunc, err = _extract_docx(raw)
        return DocumentExtractionResult(
            filename=fn,
            mime=m,
            status=st,
            text=text,
            truncated=trunc,
            error_reason=err,
        )

    if m == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        text, st, trunc, err = _extract_xlsx(raw)
        return DocumentExtractionResult(
            filename=fn,
            mime=m,
            status=st,
            text=text,
            truncated=trunc,
            error_reason=err,
        )

    if m in ("text/csv", "application/csv"):
        m = "text/csv"
        text, st, trunc, err = _extract_csv(raw)
        return DocumentExtractionResult(
            filename=fn,
            mime=m,
            status=st,
            text=text,
            truncated=trunc,
            error_reason=err,
        )

    if m == "application/vnd.ms-excel":
        return DocumentExtractionResult(
            filename=fn,
            mime=m,
            status="unsupported",
            text="",
            truncated=False,
            error_reason=None,
        )

    if m == "application/msword":
        return DocumentExtractionResult(
            filename=fn,
            mime=m,
            status="unsupported",
            text="",
            truncated=False,
            error_reason=None,
        )

    return DocumentExtractionResult(
        filename=fn,
        mime=m,
        status="unsupported",
        text="",
        truncated=False,
        error_reason=None,
    )


def _extraction_summary_line(r: DocumentExtractionResult) -> str:
    if r.status == "unsupported":
        return "unsupported for this format"
    if r.status == "failed":
        er = (r.error_reason or "error").replace("\n", " ").strip()[:120]
        return f"failed ({er})"
    if r.status == "empty":
        return "empty (no extractable text in this slice)"
    if r.status == "truncated":
        if r.mime == "application/pdf":
            return f"truncated after {MAX_EXTRACTED_CHARS_PER_FILE:,} characters or first {MAX_PDF_PAGES} PDF pages"
        if _is_spreadsheet_mime(r.mime):
            return (
                f"truncated (max {MAX_SPREADSHEET_SHEETS} sheets, {MAX_SPREADSHEET_ROWS} rows, "
                f"{MAX_SPREADSHEET_COLS} cols, or {MAX_EXTRACTED_CHARS_PER_FILE:,} chars)"
            )
        return f"truncated after {MAX_EXTRACTED_CHARS_PER_FILE:,} characters"
    if r.truncated:
        return f"extracted, truncated after {MAX_EXTRACTED_CHARS_PER_FILE:,} characters"
    return "extracted"


def format_document_block_for_llm(r: DocumentExtractionResult, *, content_body: str) -> str:
    """One document block for the model. ``content_body`` is already bounded."""
    label = "spreadsheet" if _is_spreadsheet_mime(r.mime) else "document"
    return (
        f"[Attached {label}: {r.filename}]\n"
        f"Type: {r.mime}\n"
        f"Extraction: {_extraction_summary_line(r)}\n"
        f"Content:\n{content_body}\n"
    )


def format_document_placeholder_for_llm(r: DocumentExtractionResult) -> str:
    """No extractable body: unsupported, empty, or failed — short placeholder."""
    name = r.filename
    mime = r.mime
    if r.status == "unsupported" and _is_legacy_xls_placeholder(name, mime):
        return (
            f"[Attached spreadsheet: {name}]\n"
            f"Type: {mime}\n"
            "This file was attached, but legacy .xls extraction is not enabled."
        )
    if r.status == "unsupported":
        return (
            f"[Attached document: {name}]\n"
            f"Type: {mime}\n"
            "This file was attached, but text extraction for this format is not enabled in this deployment."
        )
    if r.status == "failed":
        label = "spreadsheet" if _is_spreadsheet_mime(mime) else "document"
        return (
            f"[Attached {label}: {name}]\n"
            f"Type: {mime}\n"
            f"Extraction: {_extraction_summary_line(r)}\n"
            "Content:\n[Text extraction did not succeed; the file is still attached for your reference.]"
        )
    if r.status == "empty":
        label = "spreadsheet" if _is_spreadsheet_mime(mime) else "document"
        return (
            f"[Attached {label}: {name}]\n"
            f"Type: {mime}\n"
            f"Extraction: {_extraction_summary_line(r)}\n"
            "Content:\n[No extractable text — the document may be scanned or image-only (OCR is not enabled).]"
        )
    return format_document_block_for_llm(r, content_body=r.text)


def _skipped_budget_block(r: DocumentExtractionResult) -> str:
    label = "spreadsheet" if _is_spreadsheet_mime(r.mime) else "document"
    return (
        f"[Attached {label}: {r.filename}]\n"
        f"Type: {r.mime}\n"
        "Extraction: omitted — attached-document text budget exhausted for this message\n"
        "Content:\n[This file was not included in the extracted context due to size limits.]"
    )


def build_document_llm_sections(
    results: list[DocumentExtractionResult],
    *,
    max_total_chars: int = MAX_EXTRACTED_CHARS_PER_MESSAGE,
) -> list[str]:
    """
    Per-file sections for the LLM, enforcing ``max_total_chars`` on *extracted* body text.
    Placeholder sections (unsupported/failed/empty) do not consume the character budget.
    """
    sections: list[str] = []
    budget = max_total_chars
    for r in results:
        if r.status in ("unsupported", "failed", "empty"):
            sections.append(format_document_placeholder_for_llm(r))
            continue
        body = (r.text or "").strip()
        if not body:
            sections.append(format_document_placeholder_for_llm(r))
            continue
        if budget <= 0:
            sections.append(_skipped_budget_block(r))
            continue
        if len(body) <= budget:
            sections.append(format_document_block_for_llm(r, content_body=body))
            budget -= len(body)
            continue
        take = body[:budget]
        summary = (
            _extraction_summary_line(r)
            + f" — truncated ({MAX_EXTRACTED_CHARS_PER_MESSAGE:,} character total budget for attached documents)"
        )
        sections.append(
            f"[Attached document: {r.filename}]\n"
            f"Type: {r.mime}\n"
            f"Extraction: {summary}\n"
            f"Content:\n{take}\n"
        )
        budget = 0
    return sections

